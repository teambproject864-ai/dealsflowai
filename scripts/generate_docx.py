import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Adds subtle gray borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define border styles
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')  # Light gray
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    """Adds a heading with consistent styling and colors."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    # Custom colors
    run = p.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 139, 87)  # Sea Green
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(105, 105, 105)  # Dark Gray
        run.bold = True
    return p

def main():
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(47, 79, 79)  # Dark slate gray for readable body text

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("DealFlow.AI\nProject Documentation")
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = RGBColor(0, 128, 128)  # Teal
    run_title.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    run_sub = subtitle_p.add_run("Single Source of Truth (SSOT) — High-Level Specifications, Architecture, and Workflows")
    run_sub.font.size = Pt(12)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(105, 105, 105)

    doc.add_page_break()

    # SECTION 1: Executive Summary
    add_styled_heading(doc, "1. Executive Summary", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("1.1 Project Purpose & Overview")
    run.bold = True
    p.add_run("\nDealFlow.AI is an intelligent sales automation and pipeline enrichment platform. It leverages autonomous multi-agent reasoning, semantic memory systems, and zero-trust security guardrails to automate the lead intake, Go-To-Market (GTM) analysis, demo call scheduling, and voice confirmation pipeline. By providing instant GTM strategy generation and automated compliance-checked phone confirmations, the platform bridges the gap between digital marketing and high-converting sales calls.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("1.2 Business Value")
    run.bold = True
    p2.add_run("\n• Accelerated Sales Cycle: Eliminates manual website research by automatically scraping prospects' websites and producing customized GTM insights in under a minute.\n"
               "• Improved Demo Attendance: Increases meeting show rates by up to 24% through automated voice confirmations, text-based follow-ups, and calendar reminders.\n"
               "• Reduced Sales Overhead: Automates qualification, meeting scheduling, and post-call email recaps, allowing sales representatives to focus exclusively on closing high-probability deals.\n"
               "• SOC2-Compliant AI Interactions: Protects brand reputation and sensitive customer data using rigorous output validation and input threat detection firewalls.")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1.15
    run = p3.add_run("1.3 Key Success Metrics")
    run.bold = True
    p3.add_run("\n• Pipeline Conversion Rate: Target a 15% increase in lead-to-opportunity conversions.\n"
               "• Meeting Show Rate: Attain an 85% show rate for booked demo slots.\n"
               "• Memory OS Retrieval Latency: Keep semantic search and retrieval under 100ms via caching.\n"
               "• Compliance Adherence: 100% logs audit compliance for TCPA quiet hours and SOC2 verification scores.")

    # SECTION 2: Project Scope
    add_styled_heading(doc, "2. Project Scope", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("2.1 In-Scope Features and Capabilities")
    run.bold = True
    p.add_run("\n• Intelligent Lead Intake: Dynamic forms with real-time field validation and context-gathering questions.\n"
               "• AI-Powered GTM Analysis: Scrapes prospect domains and analyzes positioning to produce a GTM core framework, market expansion opportunities, strategy gaps, and buyer journey recommendations.\n"
               "• Autonomous Agent Brain: State-driven conversation engine supporting demo calls, live dashboard adjustments, and objections reframing.\n"
               "• Agent Security Firewall (Clawpatrol): Detects prompt injections, prevents data exfiltration (redacts PII/emails/phone numbers), computes behavior anomaly scores, and executes policy blocks.\n"
               "• Memory OS (Hermes): 4-tier persistent memory structure (working, short-term, long-term, archival) equipped with AES encryption, LRU cache, and agent-level permissions.\n"
               "• Veritas Validation Layer: Post-inference output checker that truncates responses, aligns context with company data, filters out restricted patterns, and generates compliance hashes.\n"
               "• AI Provider Router: Dynamic model and API provider routing (Hugging Face Llama, Nvidia NIMs, Kimi) based on user tier and regional requirements.\n"
               "• Multi-Agent Framework: Orchestrated data retrieval, factual validation, and report synthesis using Research, Analysis, Fact-Checking, and Synthesis agents.\n"
               "• Automated Twilio Voice Confirmation: TCPA/GDPR-compliant transactional phone alerts confirming meeting details with automated retries and parallel fallback (Email + SMS) channels.\n"
               "• Google Meet & Calendar Integration: Scheduled calendars synced with custom Google Workspace service accounts.\n"
               "• Immersive AR/3D Overlays: Interactive frontend modules highlighting pipeline performance using Framer Motion and mobile device orientation HUD.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("2.2 Out-of-Scope Elements")
    run.bold = True
    p2.add_run("\n• Google Sheets Data Synchronization: [DEPRECATED & FULLY REMOVED] Previously supported spreadsheet sync (e.g., GOOGLE_SHEET_ID, lib/sheets.ts, /api/sync/metrics) has been entirely removed from the application. All operational records (leads, calls, analyses, and audit logs) reside in Firebase Firestore.\n"
               "• Manual Outbound Calling: Sales representatives cannot make manual voice calls directly through the system.\n"
               "• Voice Stream Transfers: The virtual agent-brain does not support live transfers of active calls to human agents.")

    # SECTION 3: Core Architecture Overview
    add_styled_heading(doc, "3. Core Architecture Overview", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("3.1 Design Principles & System Components")
    run.bold = True
    p.add_run("\nDealFlow.AI is designed on a zero-trust, service-oriented architecture with decoupled layers for UI presentation, agent reasoning, data storage, and external telephony integrations.\n"
               "• Presentational Layer: Handles prospect inputs, pipeline progression, and interactive dashboards.\n"
               "• Agent Reasoning & Memory Layer: Integrates the Agent Brain with Veritas and Clawpatrol, persisting state via Hermes Memory OS.\n"
               "• Backend and Data Store Layer: Consists of Next.js API router, Firebase Firestore database, and Pinecone vector store.\n"
               "• External Services Layer: Coordinates with Twilio Voice API, Google Calendar/Meet, and Hugging Face/Nvidia/Kimi inference providers.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("3.2 Technology Stack")
    run.bold = True
    p2.add_run("\n• Frontend: Next.js (App Router), React, TailwindCSS, Framer Motion, Lucide icons.\n"
               "• Database & Vector Store: Firebase Firestore (leads, calls, analyses, audits), Pinecone Vector DB (semantic index).\n"
               "• Logic & Agent Orchestration: LangGraph (StateGraph-based workflow), Model Control Protocol (MCP).\n"
               "• Integrations: Twilio REST API (calls, SMS, webhooks), Google APIs (Calendar, Meet, Service Account), ElevenLabs API (speech synthesis).\n"
               "• Inference Providers: Hugging Face (Meta-Llama-3-8B-Instruct), Nvidia NIMs, Kimi API.")

    # SECTION 4: Key Stakeholders & Roles
    add_styled_heading(doc, "4. Key Stakeholders & Roles", level=1)
    
    # Add a table for Stakeholders
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Primary Responsibilities'
    hdr_cells[2].text = 'System Interactions / Workflows'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '008080')  # Teal header
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    stakeholders_data = [
        ("Product Owner", "Defines feature roadmap, metrics definitions, and monitors general platform performance.", "Reviews metrics dashboard, sets feature configurations."),
        ("Sales Manager", "Evaluates lead pipelines, reviews AI call summaries, and acts on forecasted deal probabilities.", "Uses Pipeline View, CallSummary component, reviews deal probability logs."),
        ("AI & Security Engineer", "Maintains Clawpatrol policies, reviews logs, adjusts prompt injection strings & anomaly models.", "Monitors Clawpatrol API endpoints, registers agent identities."),
        ("GTM Analyst", "Manages target ICP parameters, playbook documents, and updates suitability criteria functions.", "Updates suitability criteria, playbooks, review matched ICP outcomes."),
        ("Systems Admin", "Configures Google Service Accounts, Twilio credentials, and manages API resource consumption.", "Sets environment variables, manages Twilio console configuration."),
        ("Leads / Prospects", "Input profiles, browse solutions, book demo calls, receive confirmations.", "Interacts with IntakeForm, BookingWidget, receives voice alerts & SMS/Email fallbacks."),
        ("Compliance Auditors", "Inspect Firestore audit logs to verify TCPA compliance and SOC2 security verification states.", "Reviews firestore audit_logs collection, verifies Veritas complianceStatus.")
    ]
    
    for role, resp, inter in stakeholders_data:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = resp
        row_cells[2].text = inter
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 5: Primary Workflows
    add_styled_heading(doc, "5. Primary Workflows", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("5.1 Intake & GTM Analysis Workflow")
    run.bold = True
    p.add_run("\n1. Submissions: Prospect inputs company name, website, and metadata in IntakeForm.\n"
               "2. Scraping: Cheerio scraper extracts the first 4000 characters of clean text from the prospect's website.\n"
               "3. Analysis: Llama-3-8B model analyses the text and forms structured GTM analysis sections (brand overview, strategy gaps, differentiation).\n"
               "4. ICP Selection: System calculates score (0-100) and matches the best ICP playbook (Enterprise SaaS, Ecom, SMB).\n"
               "5. Presentation: Final output is rendered in LeadAnalysisDashboard and saved to Firestore.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("5.2 Booking & Voice Confirmation Workflow")
    run.bold = True
    p2.add_run("\n1. Scheduling: Prospect picks a time via BookingWidget, scheduling a Google Meet session on Google Calendar.\n"
               "2. Compliance Filter: System checks consent registry. If the user opted out, voice is bypassed. Next, checks quiet hours: if the recipient's time is outside 8 AM - 9 PM, it skips direct calling.\n"
               "3. Outbound Call: If compliant, Twilio places an outbound call reciting Alice TTS containing the booking details.\n"
               "4. Webhook tracking: status-callback route catches Twilio carrier events (ringing, completed, failed, busy).\n"
               "5. Retries & Fallback: If busy/unanswered, schedules up to 3 retries (5 min delay). If still failed, fires parallel SMS and transactional emails immediately.")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    p3.paragraph_format.line_spacing = 1.15
    run = p3.add_run("5.3 Live AI Agent Sales Call Pipeline")
    run.bold = True
    p3.add_run("\n1. Request Ingestion: Incoming prospect speech is transcribed. The prompt is inspected by Clawpatrol.\n"
               "2. Inbound Firewall: Clawpatrol scans for 15+ injection patterns and blocks if a critical threat is found.\n"
               "3. Memory Recall: System searches Hermes for historical context and stores the current prompt.\n"
               "4. Action Decision: Agent Brain reasons and returns an action (speak, show_screen, send_email) with deal probability updates.\n"
               "5. Veritas Verification: Veritas checks output length (max 5 sentences), context alignment, and compliance patterns, generating a compliance status ('SOC2_PASSED' or 'FAIL').\n"
               "6. Outbound Inspection: Clawpatrol sanitizes PII (emails/phone) before the synthesised audio response is spoken.")

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(12)
    p4.paragraph_format.line_spacing = 1.15
    run = p4.add_run("5.4 Multi-Agent Research & Synthesis Workflow")
    run.bold = True
    p4.add_run("\n1. Orchestration: Orchestrator initiates a research workflow.\n"
               "2. Message Queue: Message Queue coordinates communication across four specialized agents: Research Agent, Data Analysis Agent, Fact Checking Agent, and Synthesis Agent.\n"
               "3. Research Phase: Crawls domain and compiles factual details.\n"
               "4. Analysis Phase: Calculates descriptive metrics and checks scoring thresholds.\n"
               "5. Fact-Checking Phase: Validates claims to prevent hallucinations.\n"
               "6. Synthesis Phase: Synthesis Agent consolidates research, analysis, and validation, generating a final output report.")

    # SECTION 6: Functional Requirements
    add_styled_heading(doc, "6. Functional Requirements", level=1)
    
    # Table for functional requirements
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Priority'
    hdr_cells[3].text = 'Capability Description'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '008080')
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    func_reqs = [
        ("FR-1.1", "Core Platform", "High", "Accept prospect domain URLs, validate layout patterns, and save forms into Firestore."),
        ("FR-1.2", "Core Platform", "Medium", "Render drag-and-drop pipeline views displaying deal stages (Intake to Closed)."),
        ("FR-2.1", "AI & Automation", "High", "Scrape prospect domains, removing styling, scripts, and media elements, leaving clean content."),
        ("FR-2.2", "AI & Automation", "High", "Llama-3-8B model generates structured JSON GTM analysis reports containing strategic opportunities."),
        ("FR-2.3", "AI & Automation", "Medium", "Veritas validation layer checks context alignment, limits spoken answers to 5 sentences, and outputs audit hashes."),
        ("FR-3.1", "Outreach", "High", "Twilio automated voice confirmations placed when demo slots are booked."),
        ("FR-3.2", "Outreach", "High", "Spell out calendar codes phonetically in TTS calls to ensure maximum comprehension."),
        ("FR-3.3", "Outreach", "Medium", "Schedule and execute up to 3 retries (5-minute interval) ifTwilio voice calls fail, with automatic SMS/Email fallbacks."),
        ("FR-4.1", "Security & Compliance", "High", "Clawpatrol firewall scans inbound prompts for jailbreaks and templates, blocking critical issues."),
        ("FR-4.2", "Security & Compliance", "High", "Sanitize all PII (phone numbers, emails, cards) dynamically during inbound/outbound inspection."),
        ("FR-4.3", "Security & Compliance", "High", "Verify user consent and local time compliance (8 AM - 9 PM recipient local time) before making outbound Twilio calls.")
    ]
    
    for req_id, cat, prio, desc in func_reqs:
        row_cells = table.add_row().cells
        row_cells[0].text = req_id
        row_cells[1].text = cat
        row_cells[2].text = prio
        row_cells[3].text = desc
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 7: Non-Functional Requirements
    add_styled_heading(doc, "7. Non-Functional Requirements", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("7.1 Performance")
    run.bold = True
    p.add_run("\n• Cache Latency: In-memory LRU cache retrieval latency inside Hermes Memory OS must be under 100ms.\n"
               "• Analysis Workflows: Multi-agent synthesis and scraping must complete in under 60 seconds.\n"
               "• Telephony Webhooks: `/api/calls/status-callback` must process form payloads in under 200ms.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("7.2 Security")
    run.bold = True
    p2.add_run("\n• Encryption: All permanent and archival memories stored inside Hermes must use AES block-level encryption.\n"
               "• Access Controls: Support granular permissions checking agent-level access rules (read/write/delete) before memory retrieval.")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    p3.paragraph_format.line_spacing = 1.15
    run = p3.add_run("7.3 Reliability & Availability")
    run.bold = True
    p3.add_run("\n• Provider Redundancy: The AI Provider Router automatically fallbacks to Hugging Face if Nvidia or Kimi requests error out.\n"
               "• Telephony Resilience: If voice carrier errors occur, transaction records persist, scheduling parallel SMS/Email fallbacks.")

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(12)
    p4.paragraph_format.line_spacing = 1.15
    run = p4.add_run("7.4 Compliance")
    run.bold = True
    p4.add_run("\n• Quiet Hours Constraints: Telephony calls are blocked outside 8:00 AM - 9:00 PM local recipient time.\n"
               "• SOC2 Auditing: Immutable audit logs must write verification hashes and compliance checks for every agent interaction.")

    # SECTION 8: Project Timeline & Milestones
    add_styled_heading(doc, "8. Project Timeline & Milestones", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The development lifecycle follows a structured 5-sprint schedule, leading to security verification and production release:")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase / Sprint'
    hdr_cells[1].text = 'Deliverables'
    hdr_cells[2].text = 'Timeline'
    hdr_cells[3].text = 'Milestones'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '008080')
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    timeline_data = [
        ("Sprint 1: Foundations", "Intake forms, Pipeline Dashboard, Firestore configuration", "Weeks 1-2", "M1: Intake & Db Live"),
        ("Sprint 2: AI Analytics", "Cheerio scraping, Llama-3 integration, Multi-Agent workflow", "Weeks 3-4", "M2: GTM Engine active"),
        ("Sprint 3: Telephony", "Twilio SDK implementation, compliance checking, fallbacks", "Weeks 5-6", "M3: Voice confirmations live"),
        ("Sprint 4: Security & Storage", "Clawpatrol Firewall policies, Hermes memory encryption", "Weeks 7-8", "M4: Zero-Trust Active"),
        ("Sprint 5: Launch", "SOC2 compliance verification, system performance tuning, launch", "Weeks 9-10", "M5: Production Release")
    ]
    
    for phase, deliv, time_r, miles in timeline_data:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = deliv
        row_cells[2].text = time_r
        row_cells[3].text = miles
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 9: Risk Assessment & Mitigation
    add_styled_heading(doc, "9. Risk Assessment & Mitigation Strategies", level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Identified Risk'
    hdr_cells[1].text = 'Potential Impact'
    hdr_cells[2].text = 'Severity'
    hdr_cells[3].text = 'Mitigation Strategy'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '008080')
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    risks_data = [
        ("AI Hallucination", "Off-topic suggestions are sent to prospects, damaging professional standards.", "High", "Veritas checks output details, limits sentence size, and compares results with scraped database context."),
        ("Prompt Injection Attacks", "Jailbreaks allow bad actors to exploit the platform for unintended API tasks.", "Critical", "Clawpatrol firewall parses inputs for 15+ patterns, assigning anomaly scores and executing automatic blocks."),
        ("TCPA Violation Lawsuits", "Outbound phone calls placed to opted-out leads or during quiet hours violate compliance.", "Critical", "Compliance logic checks opt-out databases and local time windows, automatically falling back to SMS/Email."),
        ("Provider Outages", "platform fails to compile GTM analytics due to Hugging Face or Nvidia downtime.", "High", "AI Provider Router evaluates provider health and switches traffic dynamically to Kimi or Hugging Face backups.")
    ]
    
    for risk, imp, sev, mit in risks_data:
        row_cells = table.add_row().cells
        row_cells[0].text = risk
        row_cells[1].text = imp
        row_cells[2].text = sev
        row_cells[3].text = mit
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 10: Maintenance & Support Framework
    add_styled_heading(doc, "10. Maintenance & Support Framework", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("10.1 Support Tier Structure")
    run.bold = True
    p.add_run("\n• Tier 1 (Customer Operations): Solves front-end issues, registration queries, and booking adjustments.\n"
               "• Tier 2 (Integrations Support): Resolves Twilio webhook timeouts, carrier-level logs, and Google service account updates.\n"
               "• Tier 3 (AI Engineering): Tackles model router failures, modifies Clawpatrol policy conditions, and scales Hermes cache sizes.")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run("10.2 Maintenance Protocols")
    run.bold = True
    p2.add_run("\n• Weekly updates: Refresh Clawpatrol security regex lists and prompt injection signatures based on telemetry analysis.\n"
               "• Monthly audits: Adjust Hermes cache sizes and consolidation limits to scale memory speeds.\n"
               "• Quarterly checks: Audit Firestore audit logs to confirm all outbound calls adhered to local quiet hours and opt-outs.")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1.15
    run = p3.add_run("10.3 Future Roadmap")
    run.bold = True
    p3.add_run("\n• Salesforce & HubSpot CRM integration: Expand bi-directional contact syncing.\n"
               "• WhatsApp Sequencing: Build outbound templates for international prospects.\n"
               "• Advanced Search: Hybrid keyword-vector search support inside Hermes Memory OS.")

    # Save Document
    out_dir = r"d:\Project\DealFlow.AI\dealsflowsai\docs"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "PROJECT_DOCUMENTATION.docx")
    doc.save(out_path)
    print(f"Document successfully created at: {out_path}")

if __name__ == "__main__":
    main()
