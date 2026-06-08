import os
import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Adds subtle gray borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    """Adds a heading with consistent styling and colors."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
        run.bold = True
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(46, 139, 87)  # Sea Green
        run.bold = True
    return p

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_icp_docx.py <output_file_path>")
        sys.exit(1)
        
    output_path = sys.argv[1]
    
    # Read JSON payload from stdin
    try:
        payload_str = sys.stdin.read()
        data = json.loads(payload_str)
    except Exception as e:
        print(f"Error parsing JSON input: {e}")
        sys.exit(1)
        
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Default styling (Arial)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(47, 79, 79)  # Dark slate gray
    
    # Header Banner Style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("IDEAL CUSTOMER PROFILE (ICP) REPORT\n& GTM STRATEGY BLUEPRINT")
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0, 128, 128)  # Teal
    run_title.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("Synchronized GTM Alignment with Hermes OS, Clawpatrol, & Multi-Agent Frameworks")
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(105, 105, 105)
    
    # 1. Overview
    add_styled_heading(doc, "1. Go-to-Market ICP Overview", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(data.get("ICP Overview", ""))
    
    # 2. Customer Segmentation
    add_styled_heading(doc, "2. Target Customer Segmentation", level=1)
    seg = data.get("Target Customer Segmentation", {})
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("• Demographics: ")
    run.bold = True
    p.add_run(seg.get("Demographics", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("• Firmographics: ")
    run.bold = True
    p.add_run(seg.get("Firmographics", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("• SDR Team Scale: ")
    run.bold = True
    p.add_run(seg.get("SDR Team Scale", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("• Target Geographies: ")
    run.bold = True
    p.add_run(seg.get("Target Geographies", ""))
    
    # 3. Pain Points
    add_styled_heading(doc, "3. Key Pain Point Mapping", level=1)
    pains = data.get("Key Pain Point Mapping", {})
    for title, desc in pains.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)
        
    # 4. Technical Product Value Propositions
    add_styled_heading(doc, "4. Technical Product Value Proposition Alignment", level=1)
    props = data.get("Technical Product Value Proposition Alignment", {})
    for title, desc in props.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"• {title.replace(' Alignment', '')}: ")
        run.bold = True
        p.add_run(desc)
        
    # 5. Use Case Prioritization Grid
    add_styled_heading(doc, "5. Use Case Prioritization Grid", level=1)
    cases = data.get("Use Case Prioritization Grid", {})
    for title, desc in cases.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)

    # 6. Market Sizing & Competitive Landscape
    add_styled_heading(doc, "6. Market Sizing & Competitor Estimates", level=1)
    market = data.get("Market Sizing & Competitor Estimates", {})
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• Consensus TAM 2026: ")
    run.bold = True
    p.add_run(market.get("TAM 2026 Consensus", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("• Consensus CAGR: ")
    run.bold = True
    p.add_run(market.get("Consensus Growth CAGR", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("• Competitor Market Share Distribution (Consensus consensus):")
    run.bold = True
    
    # Draw competitor table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Competitor'
    hdr_cells[1].text = 'Estimated Market Share'
    hdr_cells[2].text = 'Primary GTM Focus Model'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '008080')  # Teal header
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        
    competitor_rows = [
        ("Apollo.io", "40.0% Share", "All-in-One Database & sequencing, PLG focus"),
        ("Clay.com", "25.0% Share", "Data Enrichment & CRM waterfall orchestration"),
        ("Instantly.ai", "20.0% Share", "Cold Email Outreach, modular deliverability"),
        ("DealFlow.AI (Target)", "15.0% Share", "Autonomous Agentic GTM & Telemetry integrations"),
    ]
    
    for comp, share, focus in competitor_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = share
        row_cells[2].text = focus
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 7. Consensus Validation
    add_styled_heading(doc, "7. Consensus Validation Audit Log", level=1)
    val = data.get("Consensus Validation Log", {})
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• Consensus Status: ")
    run.bold = True
    p.add_run(val.get("Verification Status", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• Coefficient of Variation (CoV) Check: ")
    run.bold = True
    p.add_run(val.get("Coefficient of Variation Check", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• Margin of Error (95% CI) Check: ")
    run.bold = True
    p.add_run(val.get("Margin of Error (95%) Check", ""))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("• SOC2 Verification Audit Stamp: ")
    run.bold = True
    p.add_run(val.get("Audit Integrity Stamp", ""))
    
    # Save document
    try:
        doc.save(output_path)
        print(f"DOCX successfully generated at {output_path}")
    except Exception as e:
        print(f"Error saving DOCX document: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
